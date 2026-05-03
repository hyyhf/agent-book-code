"""
FunHarness - Session attachment management and document parsing.

Attachments are copied into .funharness/uploads/<session_id>/ so the agent can
refer to stable, workspace-local paths instead of arbitrary user paths.
"""
from __future__ import annotations

import csv
import hashlib
import json
import mimetypes
import shutil
import uuid
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path


ATTACHMENT_PREVIEW_CHARS = 1000
DEFAULT_ATTACHMENT_MAX_CHARS = 120_000
TEXT_EXTENSIONS = {
    ".txt", ".md", ".markdown", ".py", ".js", ".ts", ".tsx", ".jsx",
    ".json", ".yaml", ".yml", ".toml", ".ini", ".cfg", ".csv", ".tsv",
    ".html", ".css", ".xml", ".sql", ".sh", ".bat", ".ps1", ".log",
    ".rst", ".tex", ".java", ".c", ".cpp", ".h", ".hpp", ".go", ".rs",
    ".rb", ".php",
}
UNSUPPORTED_OFFICE_EXTENSIONS = {".doc", ".xls"}


@dataclass
class AttachmentRecord:
    id: str
    original_name: str
    stored_path: str
    extension: str
    mime_type: str
    size: int
    added_at: str
    parse_status: str
    preview: str

    @classmethod
    def from_dict(cls, data: dict) -> "AttachmentRecord":
        return cls(
            id=data["id"],
            original_name=data.get("original_name", ""),
            stored_path=data.get("stored_path", ""),
            extension=data.get("extension", ""),
            mime_type=data.get("mime_type", ""),
            size=int(data.get("size", 0)),
            added_at=data.get("added_at", ""),
            parse_status=data.get("parse_status", ""),
            preview=data.get("preview", ""),
        )


def _truncate(text: str, max_chars: int) -> str:
    if max_chars < 0:
        max_chars = 0
    if len(text) <= max_chars:
        return text
    return text[:max_chars].rstrip() + f"\n...(truncated, total {len(text)} chars)"


def parse_document(path: str | Path, max_chars: int = DEFAULT_ATTACHMENT_MAX_CHARS) -> str:
    """Extract readable text from common attachment formats."""
    file_path = Path(path)
    ext = file_path.suffix.lower()

    if ext in UNSUPPORTED_OFFICE_EXTENSIONS:
        return (
            f"Unsupported legacy Office format: {ext}. "
            "Please convert the file to .docx or .xlsx and attach it again."
        )
    if ext == ".pdf":
        text = _parse_pdf(file_path)
    elif ext == ".docx":
        text = _parse_docx(file_path)
    elif ext == ".xlsx":
        text = _parse_xlsx(file_path)
    elif ext in {".csv", ".tsv"}:
        text = _parse_delimited(file_path, delimiter="\t" if ext == ".tsv" else ",")
    elif ext in TEXT_EXTENSIONS or _looks_like_text(file_path):
        text = _parse_text(file_path)
    else:
        return (
            f"Unsupported or binary file format: {ext or '(no extension)'}. "
            "Supported attachment formats include text, PDF, DOCX, XLSX, CSV, and TSV."
        )

    return _truncate(text, max_chars)


def _parse_text(path: Path) -> str:
    text = path.read_text(encoding="utf-8", errors="replace")
    line_count = len(text.splitlines()) if text else 0
    return f"[{line_count} lines]\n{text}"


def _parse_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return "PDF parsing requires the 'pypdf' package."

    try:
        reader = PdfReader(str(path))
    except Exception as exc:
        return f"PDF parse failed: {exc}"

    pages = []
    for index, page in enumerate(reader.pages, 1):
        try:
            page_text = page.extract_text() or ""
        except Exception as exc:
            page_text = f"[page extraction failed: {exc}]"
        page_text = page_text.strip()
        if page_text:
            pages.append(f"--- Page {index} ---\n{page_text}")

    if not pages:
        return f"[PDF: {len(reader.pages)} pages]\n(no extractable text)"
    return f"[PDF: {len(reader.pages)} pages]\n" + "\n\n".join(pages)


def _parse_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        return "DOCX parsing requires the 'python-docx' package."

    try:
        doc = Document(str(path))
    except Exception as exc:
        return f"DOCX parse failed: {exc}"

    parts = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table_index, table in enumerate(doc.tables, 1):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            parts.append(f"[Table {table_index}]\n" + "\n".join(rows))
    return "[DOCX]\n" + ("\n\n".join(parts) if parts else "(no text)")


def _parse_xlsx(path: Path) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return "XLSX parsing requires the 'openpyxl' package."

    try:
        workbook = load_workbook(str(path), read_only=True, data_only=True)
    except Exception as exc:
        return f"XLSX parse failed: {exc}"

    parts = []
    try:
        for sheet in workbook.worksheets:
            rows = []
            for row in sheet.iter_rows(values_only=True):
                values = ["" if value is None else str(value) for value in row]
                if any(value.strip() for value in values):
                    rows.append("\t".join(values).rstrip())
            bounds = f"{sheet.max_row} rows x {sheet.max_column} cols"
            body = "\n".join(rows) if rows else "(empty sheet)"
            parts.append(f"--- Sheet: {sheet.title} ({bounds}) ---\n{body}")
    finally:
        workbook.close()
    return "[XLSX]\n" + "\n\n".join(parts)


def _parse_delimited(path: Path, delimiter: str) -> str:
    text = path.read_text(encoding="utf-8", errors="replace")
    rows = []
    reader = csv.reader(text.splitlines(), delimiter=delimiter)
    for row in reader:
        rows.append("\t".join(row))
    label = "TSV" if delimiter == "\t" else "CSV"
    return f"[{label}: {len(rows)} rows]\n" + "\n".join(rows)


def _looks_like_text(path: Path) -> bool:
    try:
        chunk = path.read_bytes()[:4096]
    except OSError:
        return False
    if b"\x00" in chunk:
        return False
    try:
        chunk.decode("utf-8")
        return True
    except UnicodeDecodeError:
        return False


class AttachmentManager:
    """Persistent attachment list for one conversation session."""

    def __init__(self, session_id: str, root: str | Path = ".funharness/uploads"):
        self.session_id = session_id
        self.root = Path(root)
        self.session_dir = self.root / session_id
        self.index_path = self.session_dir / "attachments.json"
        self.records: list[AttachmentRecord] = []
        self.load()

    def load(self) -> None:
        if not self.index_path.exists():
            self.records = []
            return
        data = json.loads(self.index_path.read_text(encoding="utf-8"))
        self.records = [AttachmentRecord.from_dict(item) for item in data.get("attachments", [])]

    def save(self) -> None:
        self.session_dir.mkdir(parents=True, exist_ok=True)
        payload = {"session_id": self.session_id, "attachments": [asdict(r) for r in self.records]}
        self.index_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    def add(self, source_path: str | Path) -> AttachmentRecord:
        source = Path(source_path).expanduser()
        if not source.is_absolute():
            source = Path.cwd() / source
        source = source.resolve()
        if not source.is_file():
            raise FileNotFoundError(f"Attachment file not found: {source_path}")

        self.session_dir.mkdir(parents=True, exist_ok=True)
        attachment_id = self._new_id(source)
        stored_name = f"{attachment_id}{source.suffix.lower()}"
        stored_path = self.session_dir / stored_name
        shutil.copy2(source, stored_path)

        mime_type = mimetypes.guess_type(source.name)[0] or "application/octet-stream"
        preview_text = parse_document(stored_path, max_chars=ATTACHMENT_PREVIEW_CHARS)
        parse_status = "ok"
        if preview_text.startswith(("Unsupported", "PDF parse failed", "DOCX parse failed", "XLSX parse failed")):
            parse_status = "unsupported" if preview_text.startswith("Unsupported") else "error"

        record = AttachmentRecord(
            id=attachment_id,
            original_name=source.name,
            stored_path=str(stored_path),
            extension=source.suffix.lower(),
            mime_type=mime_type,
            size=stored_path.stat().st_size,
            added_at=datetime.now().isoformat(timespec="seconds"),
            parse_status=parse_status,
            preview=preview_text,
        )
        self.records.append(record)
        self.save()
        return record

    def list(self) -> list[AttachmentRecord]:
        return list(self.records)

    def get(self, attachment_id: str) -> AttachmentRecord | None:
        for record in self.records:
            if record.id == attachment_id:
                return record
        return None

    def detach(self, target: str) -> str:
        if target.lower() == "all":
            count = len(self.records)
            self.records = []
            self.save()
            return f"Detached {count} attachment(s). Files remain on disk."

        before = len(self.records)
        self.records = [record for record in self.records if record.id != target]
        if len(self.records) == before:
            return f"Attachment not found: {target}"
        self.save()
        return f"Detached attachment: {target}. File remains on disk."

    def summary(self, include_preview: bool = False) -> str:
        if not self.records:
            return "(no attachments)"
        lines = ["[Attached Files]"]
        for record in self.records:
            lines.append(
                f"- id={record.id} name={record.original_name} "
                f"type={record.extension or record.mime_type} size={record.size} "
                f"path={record.stored_path} status={record.parse_status}"
            )
            if include_preview and record.preview:
                lines.append("  preview:")
                for preview_line in record.preview.splitlines()[:12]:
                    lines.append(f"    {preview_line}")
        return "\n".join(lines)

    def read(self, attachment_id: str, max_chars: int = DEFAULT_ATTACHMENT_MAX_CHARS) -> str:
        record = self.get(attachment_id)
        if record is None:
            return f"Attachment not found: {attachment_id}"
        return parse_document(record.stored_path, max_chars=max_chars)

    @staticmethod
    def _new_id(path: Path) -> str:
        seed = f"{path.name}:{path.stat().st_size}:{uuid.uuid4().hex}"
        return "att_" + hashlib.sha1(seed.encode("utf-8")).hexdigest()[:10]
