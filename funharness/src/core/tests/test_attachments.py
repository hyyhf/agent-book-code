from __future__ import annotations

import os
import tempfile
import unittest
from pathlib import Path

from funharness.src.core.attachments import AttachmentManager, parse_document


class AttachmentParsingTests(unittest.TestCase):
    def test_reads_text_and_truncates(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "notes.txt"
            path.write_text("alpha\nbeta\ngamma\n", encoding="utf-8")

            result = parse_document(path, max_chars=20)

        self.assertIn("[3 lines]", result)
        self.assertIn("alpha", result)
        self.assertIn("truncated", result)

    def test_extracts_docx_paragraphs_and_tables(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "report.docx"
            doc = Document()
            doc.add_paragraph("Executive summary")
            table = doc.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "Metric"
            table.cell(0, 1).text = "Value"
            doc.save(path)

            result = parse_document(path)

        self.assertIn("[DOCX]", result)
        self.assertIn("Executive summary", result)
        self.assertIn("Metric | Value", result)

    def test_extracts_xlsx_multiple_sheets(self) -> None:
        from openpyxl import Workbook

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "book.xlsx"
            workbook = Workbook()
            first = workbook.active
            first.title = "Sales"
            first.append(["Region", "Revenue"])
            first.append(["East", 42])
            second = workbook.create_sheet("Notes")
            second.append(["hello"])
            workbook.save(path)

            result = parse_document(path)

        self.assertIn("--- Sheet: Sales", result)
        self.assertIn("Region\tRevenue", result)
        self.assertIn("--- Sheet: Notes", result)
        self.assertIn("hello", result)

    def test_extracts_text_pdf_and_handles_empty_pdf(self) -> None:
        from pypdf import PdfWriter

        with tempfile.TemporaryDirectory() as tmp:
            text_pdf = Path(tmp) / "text.pdf"
            empty_pdf = Path(tmp) / "empty.pdf"
            _write_simple_text_pdf(text_pdf, "Hello PDF")
            writer = PdfWriter()
            writer.add_blank_page(width=72, height=72)
            with empty_pdf.open("wb") as handle:
                writer.write(handle)

            text_result = parse_document(text_pdf)
            empty_result = parse_document(empty_pdf)

        self.assertIn("Hello PDF", text_result)
        self.assertIn("no extractable text", empty_result)

    def test_legacy_office_formats_are_unsupported(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "old.doc"
            path.write_bytes(b"legacy")

            result = parse_document(path)

        self.assertIn("Unsupported legacy Office format", result)


class AttachmentManagerTests(unittest.TestCase):
    def test_adds_multiple_files_and_persists_index(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            first = root / "a.txt"
            second = root / "b.txt"
            first.write_text("one", encoding="utf-8")
            second.write_text("two", encoding="utf-8")

            manager = AttachmentManager("session1", root=root / "uploads")
            record1 = manager.add(first)
            record2 = manager.add(second)
            restored = AttachmentManager("session1", root=root / "uploads")

        self.assertNotEqual(record1.id, record2.id)
        self.assertEqual(len(restored.list()), 2)
        self.assertTrue(Path(record1.stored_path).name.startswith(record1.id))

    def test_detach_removes_references_without_deleting_files(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "a.txt"
            source.write_text("one", encoding="utf-8")

            manager = AttachmentManager("session1", root=root / "uploads")
            record = manager.add(source)
            stored = Path(record.stored_path)
            message = manager.detach(record.id)
            all_message = manager.detach("all")

            self.assertIn("Detached attachment", message)
            self.assertIn("Detached 0 attachment", all_message)
            self.assertEqual(manager.list(), [])
            self.assertTrue(stored.exists())

    def test_relative_attach_paths_resolve_from_cwd(self) -> None:
        previous_cwd = Path.cwd()
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "local.txt"
            source.write_text("local", encoding="utf-8")
            try:
                os.chdir(root)
                manager = AttachmentManager("session1", root=root / "uploads")
                record = manager.add("local.txt")
            finally:
                os.chdir(previous_cwd)

        self.assertEqual(record.original_name, "local.txt")


def _write_simple_text_pdf(path: Path, text: str) -> None:
    safe_text = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        f"<< /Length {len(f'BT /F1 12 Tf 72 720 Td ({safe_text}) Tj ET'.encode('ascii'))} >>\n"
        f"stream\nBT /F1 12 Tf 72 720 Td ({safe_text}) Tj ET\nendstream".encode("ascii"),
    ]
    output = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, 1):
        offsets.append(len(output))
        output.extend(f"{index} 0 obj\n".encode("ascii"))
        output.extend(obj)
        output.extend(b"\nendobj\n")
    xref = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    output.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref}\n%%EOF\n".encode("ascii")
    )
    path.write_bytes(bytes(output))


if __name__ == "__main__":
    unittest.main()
